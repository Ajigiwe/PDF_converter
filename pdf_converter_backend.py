#!/usr/bin/env python3
"""
PDF Converter & Merger Backend
Handles conversions between PDF, DOCX, and Markdown formats
Features: file conversion, merging, ASCII art to diagram conversion, custom themes, headers/footers
"""

import os
import re
import io
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

# PDF handling
from pdf2image import convert_from_path
from PyPDF2 import PdfWriter, PdfReader
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# Document handling
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Image processing
from PIL import Image, ImageDraw
import markdown


@dataclass
class ConversionOptions:
    """Configuration for document conversions"""
    output_format: str = 'pdf'  # pdf, docx, md
    theme: str = 'light'  # light, dark, professional, minimal
    add_header: bool = True
    header_text: str = 'Document'
    footer_text: str = ''
    font_size: int = 12
    margins: float = 1.0  # in inches
    page_size: str = 'letter'  # letter, a4
    add_page_numbers: bool = True
    convert_ascii_diagrams: bool = True
    pages_per_sheet: int = 1


class ThemeConfig:
    """Theme definitions for document styling"""
    
    THEMES = {
        'light': {
            'bg_color': (255, 255, 255),
            'text_color': (0, 0, 0),
            'accent_color': (41, 128, 185),
            'header_bg': (236, 240, 241),
            'font_family': 'Helvetica',
        },
        'dark': {
            'bg_color': (44, 47, 51),
            'text_color': (236, 240, 241),
            'accent_color': (52, 152, 219),
            'header_bg': (52, 73, 94),
            'font_family': 'Helvetica',
        },
        'professional': {
            'bg_color': (255, 255, 255),
            'text_color': (33, 33, 33),
            'accent_color': (0, 51, 102),
            'header_bg': (230, 230, 230),
            'font_family': 'Times-Roman',
        },
        'minimal': {
            'bg_color': (255, 255, 255),
            'text_color': (50, 50, 50),
            'accent_color': (100, 100, 100),
            'header_bg': (245, 245, 245),
            'font_family': 'Helvetica',
        }
    }

    @staticmethod
    def get_theme(theme_name: str) -> Dict:
        return ThemeConfig.THEMES.get(theme_name, ThemeConfig.THEMES['light'])


class ASCIIDiagramConverter:
    """Convert ASCII art diagrams to proper SVG/PNG images"""
    
    @staticmethod
    def detect_ascii_diagrams(text: str) -> List[Tuple[int, int, str]]:
        """Detect ASCII diagram blocks in text"""
        lines = text.split('\n')
        diagrams = []
        in_diagram = False
        start_idx = 0
        diagram_lines = []
        
        for i, line in enumerate(lines):
            # Detect diagram markers or ASCII art patterns
            if re.match(r'^\s*[\+\-\|\/\\]+', line) or re.match(r'^\s*[\[\{\(<>)}\]]+', line):
                if not in_diagram:
                    in_diagram = True
                    start_idx = i
                diagram_lines.append(line)
            else:
                if in_diagram and diagram_lines:
                    diagrams.append((start_idx, i, '\n'.join(diagram_lines)))
                    in_diagram = False
                    diagram_lines = []
        
        if in_diagram and diagram_lines:
            diagrams.append((start_idx, len(lines), '\n'.join(diagram_lines)))
        
        return diagrams
    
    @staticmethod
    def convert_to_svg(ascii_art: str) -> str:
        """Convert ASCII art to SVG (basic implementation)"""
        lines = ascii_art.split('\n')
        width = max(len(line) for line in lines) if lines else 0
        height = len(lines)
        
        # Simple SVG generation
        svg_width = width * 8
        svg_height = height * 16
        
        svg = f'''<svg width="{svg_width}" height="{svg_height}" xmlns="http://www.w3.org/2000/svg">
            <rect width="{svg_width}" height="{svg_height}" fill="white" stroke="black" stroke-width="1"/>
            <g font-family="monospace" font-size="12">'''
        
        for y, line in enumerate(lines):
            svg += f'\n<text x="4" y="{y * 16 + 14}" fill="black">{line}</text>'
        
        svg += '\n</g>\n</svg>'
        return svg


class DocumentConverter:
    """Main converter class for handling file conversions"""
    
    def __init__(self, options: ConversionOptions = None):
        self.options = options or ConversionOptions()
        self.theme = ThemeConfig.get_theme(self.options.theme)
    
    def convert_file(self, input_path: str, output_path: str) -> bool:
        """
        Convert a file based on input and output extensions
        Returns: True if successful, False otherwise
        """
        input_ext = Path(input_path).suffix.lower()
        output_ext = Path(output_path).suffix.lower()
        
        try:
            # Determine conversion path
            if output_ext == '.pdf':
                return self._convert_to_pdf(input_path, output_path, input_ext)
            elif output_ext == '.docx':
                return self._convert_to_docx(input_path, output_path, input_ext)
            elif output_ext == '.md':
                return self._convert_to_markdown(input_path, output_path, input_ext)
            else:
                raise ValueError(f"Unsupported output format: {output_ext}")
        
        except Exception as e:
            print(f"Conversion error: {str(e)}")
            return False
    
    def _convert_to_pdf(self, input_path: str, output_path: str, input_ext: str) -> bool:
        """Convert document to PDF format"""
        if input_ext == '.pdf':
            # PDF to PDF (copy with processing)
            self._apply_pdf_enhancements(input_path, output_path)
            return True
        elif input_ext == '.docx':
            return self._docx_to_pdf(input_path, output_path)
        elif input_ext == '.md':
            return self._markdown_to_pdf(input_path, output_path)
        elif input_ext == '.txt':
            return self._text_to_pdf(input_path, output_path)
        
        return False
    
    def _convert_to_docx(self, input_path: str, output_path: str, input_ext: str) -> bool:
        """Convert document to DOCX format"""
        if input_ext == '.docx':
            # DOCX to DOCX (copy)
            import shutil
            shutil.copy(input_path, output_path)
            return True
        elif input_ext == '.pdf':
            return self._pdf_to_docx(input_path, output_path)
        elif input_ext == '.md':
            return self._markdown_to_docx(input_path, output_path)
        elif input_ext == '.txt':
            return self._text_to_docx(input_path, output_path)
        
        return False
    
    def _convert_to_markdown(self, input_path: str, output_path: str, input_ext: str) -> bool:
        """Convert document to Markdown format"""
        if input_ext == '.md':
            # MD to MD (copy)
            import shutil
            shutil.copy(input_path, output_path)
            return True
        elif input_ext == '.txt':
            return self._text_to_markdown(input_path, output_path)
        
        # For PDF and DOCX, we'd need OCR/complex extraction
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# Document\n\n{content}\n")
        
        return True
    
    def _docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """Convert DOCX to PDF using reportlab"""
        try:
            doc = Document(docx_path)
            elements = []
            
            # Add header
            if self.options.add_header:
                elements.append(self._create_header_para())
                elements.append(Spacer(1, 0.3 * inch))
            
            # Extract paragraphs and tables
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    para = Paragraph(
                        element.text or '&nbsp;',
                        getSampleStyleSheet()['Normal']
                    )
                    elements.append(para)
                elif element.tag.endswith('tbl'):
                    # Handle tables
                    pass
            
            # Create PDF
            doc_pdf = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                topMargin=self.options.margins * inch,
                bottomMargin=self.options.margins * inch,
                leftMargin=self.options.margins * inch,
                rightMargin=self.options.margins * inch,
            )
            
            doc_pdf.build(elements)
            return True
        except Exception as e:
            print(f"DOCX to PDF error: {str(e)}")
            return False
    
    def _text_to_pdf(self, text_path: str, pdf_path: str) -> bool:
        """Convert plain text to PDF"""
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            elements = []
            
            if self.options.add_header:
                elements.append(self._create_header_para())
                elements.append(Spacer(1, 0.3 * inch))
            
            # Split content into paragraphs
            for line in content.split('\n'):
                if line.strip():
                    style = getSampleStyleSheet()['Normal']
                    style.fontSize = self.options.font_size
                    elements.append(Paragraph(line, style))
                else:
                    elements.append(Spacer(1, 0.1 * inch))
            
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                topMargin=self.options.margins * inch,
                bottomMargin=self.options.margins * inch,
            )
            doc.build(elements)
            return True
        except Exception as e:
            print(f"Text to PDF error: {str(e)}")
            return False
    
    def _markdown_to_pdf(self, md_path: str, pdf_path: str) -> bool:
        """Convert Markdown to PDF"""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert markdown to HTML
            html_content = markdown.markdown(md_content)
            
            # Create PDF with styled content
            elements = []
            
            if self.options.add_header:
                elements.append(self._create_header_para())
                elements.append(Spacer(1, 0.3 * inch))
            
            # Parse HTML and convert to reportlab elements
            style_sheet = getSampleStyleSheet()
            for line in html_content.split('\n'):
                if line.strip():
                    if line.startswith('<h1>'):
                        text = re.sub(r'<[^>]+>', '', line)
                        elements.append(Paragraph(text, style_sheet['Heading1']))
                    elif line.startswith('<h2>'):
                        text = re.sub(r'<[^>]+>', '', line)
                        elements.append(Paragraph(text, style_sheet['Heading2']))
                    elif line.startswith('<p>'):
                        text = re.sub(r'<[^>]+>', '', line)
                        elements.append(Paragraph(text, style_sheet['Normal']))
            
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                topMargin=self.options.margins * inch,
                bottomMargin=self.options.margins * inch,
            )
            doc.build(elements)
            return True
        except Exception as e:
            print(f"Markdown to PDF error: {str(e)}")
            return False
    
    def _text_to_docx(self, text_path: str, docx_path: str) -> bool:
        """Convert plain text to DOCX"""
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document()
            
            if self.options.add_header:
                doc.add_heading(self.options.header_text, level=1)
            
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            doc.save(docx_path)
            return True
        except Exception as e:
            print(f"Text to DOCX error: {str(e)}")
            return False
    
    def _markdown_to_docx(self, md_path: str, docx_path: str) -> bool:
        """Convert Markdown to DOCX"""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document()
            
            current_style = 'Normal'
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
            
            doc.save(docx_path)
            return True
        except Exception as e:
            print(f"Markdown to DOCX error: {str(e)}")
            return False
    
    def _text_to_markdown(self, text_path: str, md_path: str) -> bool:
        """Convert plain text to Markdown"""
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(f"# Document\n\n{content}\n")
            
            return True
        except Exception as e:
            print(f"Text to Markdown error: {str(e)}")
            return False
    
    def _pdf_to_docx(self, pdf_path: str, docx_path: str) -> bool:
        """Convert PDF to DOCX (requires pdfminer or similar)"""
        try:
            doc = Document()
            
            if self.options.add_header:
                doc.add_heading(self.options.header_text, level=1)
            
            # Basic PDF text extraction
            with open(pdf_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line)
            
            doc.save(docx_path)
            return True
        except Exception as e:
            print(f"PDF to DOCX error: {str(e)}")
            return False
    
    def _apply_pdf_enhancements(self, input_pdf: str, output_pdf: str) -> bool:
        """Apply theme and header enhancements to PDF"""
        try:
            reader = PdfReader(input_pdf)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            with open(output_pdf, 'wb') as f:
                writer.write(f)
            
            return True
        except Exception as e:
            print(f"PDF enhancement error: {str(e)}")
            return False
    
    def _create_header_para(self) -> Paragraph:
        """Create a styled header paragraph"""
        style = ParagraphStyle(
            'CustomHeader',
            parent=getSampleStyleSheet()['Heading1'],
            fontSize=16,
            textColor=colors.HexColor(f"#{self.theme['accent_color'][0]:02x}{self.theme['accent_color'][1]:02x}{self.theme['accent_color'][2]:02x}"),
            spaceAfter=12,
            alignment=0,
        )
        return Paragraph(self.options.header_text, style)


class PDFMerger:
    """Merge multiple PDF, DOCX, and other document files"""
    
    @staticmethod
    def merge_pdfs(pdf_files: List[str], output_path: str, options: ConversionOptions = None) -> bool:
        """Merge multiple PDF files into one"""
        try:
            options = options or ConversionOptions()
            n = options.pages_per_sheet
            
            writer = PdfWriter()
            
            all_pages = []
            for pdf_file in pdf_files:
                reader = PdfReader(pdf_file)
                all_pages.extend(reader.pages)
                
            num_pages = len(all_pages)
            if num_pages == 0:
                return False
                
            if n <= 1:
                for page in all_pages:
                    writer.add_page(page)
            else:
                from PyPDF2 import PageObject, Transformation
                
                ref_page = all_pages[0]
                out_width = float(ref_page.mediabox.width)
                out_height = float(ref_page.mediabox.height)
                
                if n == 2:
                    cols, rows = 1, 2
                elif n == 4:
                    cols, rows = 2, 2
                elif n == 6:
                    cols, rows = 2, 3
                elif n == 9:
                    cols, rows = 3, 3
                else:
                    cols, rows = 2, 2
                    n = 4
                    
                scale_x = out_width / cols / out_width
                scale_y = out_height / rows / out_height
                scale = min(scale_x, scale_y) * 0.95 # 5% padding
                
                cell_width = out_width / cols
                cell_height = out_height / rows
                
                for i in range(0, num_pages, n):
                    out_page = PageObject.create_blank_page(width=out_width, height=out_height)
                    for j in range(n):
                        if i + j < num_pages:
                            src_page = all_pages[i + j]
                            col = j % cols
                            row = j // cols
                            
                            tx = col * cell_width + (cell_width - float(src_page.mediabox.width) * scale) / 2
                            ty = out_height - (row + 1) * cell_height + (cell_height - float(src_page.mediabox.height) * scale) / 2
                            
                            op = Transformation().scale(sx=scale, sy=scale).translate(tx=tx, ty=ty)
                            src_page.add_transformation(op)
                            out_page.merge_page(src_page)
                    writer.add_page(out_page)
            
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            return True
        except Exception as e:
            print(f"PDF merge error: {str(e)}")
            return False
    
    @staticmethod
    def merge_to_docx(files: List[str], output_path: str, options: ConversionOptions = None) -> bool:
        """Merge multiple documents into a DOCX file"""
        try:
            options = options or ConversionOptions()
            doc = Document()
            
            if options.add_header:
                doc.add_heading(options.header_text, level=1)
            
            for file_path in files:
                ext = Path(file_path).suffix.lower()
                
                # Add page break between documents
                doc.add_page_break()
                doc.add_heading(Path(file_path).stem, level=2)
                
                if ext == '.docx':
                    src_doc = Document(file_path)
                    for element in src_doc.element.body:
                        doc.element.body.append(element._element)
                
                elif ext == '.txt' or ext == '.md':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    for line in content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
                
                elif ext == '.pdf':
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            for line in text.split('\n'):
                                if line.strip():
                                    doc.add_paragraph(line)
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Document merge error: {str(e)}")
            return False


if __name__ == '__main__':
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="PDF Converter Backend")
    parser.add_argument('--merge', action='store_true', help="Merge mode")
    parser.add_argument('--input', type=str, help="Input file path for convert")
    parser.add_argument('--inputs', type=str, help="Comma-separated input paths for merge")
    parser.add_argument('--output', type=str, required=True, help="Output file path")
    parser.add_argument('--format', type=str, default='pdf', help="Output format")
    parser.add_argument('--theme', type=str, default='light', help="Theme")
    parser.add_argument('--add-header', type=str, default='1', help="Add header flag")
    parser.add_argument('--header-text', type=str, default='Document', help="Header text")
    parser.add_argument('--font-size', type=int, default=12, help="Font size")
    parser.add_argument('--margins', type=float, default=1.0, help="Margins")
    parser.add_argument('--pages-per-sheet', type=int, default=1, help="Pages per sheet")

    # Sometimes args have quotes due to shell escaping, let's strip them
    def strip_quotes(s):
        if not s: return s
        if s.startswith('"') and s.endswith('"'): return s[1:-1]
        if s.startswith("'") and s.endswith("'"): return s[1:-1]
        return s

    args = parser.parse_args()

    options = ConversionOptions(
        output_format=args.format,
        theme=args.theme,
        add_header=(args.add_header == '1' or args.add_header.lower() == 'true'),
        header_text=strip_quotes(args.header_text),
        font_size=args.font_size,
        margins=args.margins,
        pages_per_sheet=args.pages_per_sheet
    )

    output_path = strip_quotes(args.output)

    if args.merge:
        if not args.inputs:
            print("Error: --inputs required for merge")
            sys.exit(1)
        
        inputs = [strip_quotes(p) for p in args.inputs.split(',')]
        
        merger = PDFMerger()
        if args.format.lower() == 'pdf':
            success = merger.merge_pdfs(inputs, output_path, options)
        else:
            success = merger.merge_to_docx(inputs, output_path, options)
            
        if not success:
            sys.exit(1)
    else:
        if not args.input:
            print("Error: --input required for convert")
            sys.exit(1)
            
        input_path = strip_quotes(args.input)
        
        converter = DocumentConverter(options)
        success = converter.convert_file(input_path, output_path)
        
        if not success:
            sys.exit(1)
    
    sys.exit(0)
